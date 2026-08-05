"""
SENTRY — Cyber Security Threat Intelligence Assistant (Enterprise SOC Edition)

Single-file build for Streamlit Community Cloud deployment.

Deploy on Streamlit Community Cloud (share.streamlit.io):
    1. Push this app.py + requirements.txt to a public GitHub repo.
    2. On streamlit.io, "New app" -> point to the repo -> main file: app.py
    3. In App settings -> Secrets, add:
         GEMINI_API_KEY = "your-gemini-key"   # preferred LLM — https://aistudio.google.com/apikey
         GROQ_API_KEY   = "your-groq-key"     # fallback LLM  — https://console.groq.com/keys
         NVD_API_KEY    = "your-nvd-key"      # optional, raises NVD rate limits

Everything below is organized into clearly labeled sections (all originally
separate modules) so it's still easy to navigate despite being one file:
  1. Config              5. Confidence scoring   9. Report generator (PDF/DOCX)
  2. LLM client           6. Ingestion (multi-fmt) 10. Local storage (history/reports)
  3. Vector store (Chroma) 7. Prompts              11. Streamlit UI
  4. NVD + CISA KEV client  8. LangGraph workflow
"""
from __future__ import annotations


# ============================================================================
# SECTION: core/config.py
# ============================================================================

import os
from dataclasses import dataclass, field
from pathlib import Path

try:
    import streamlit as st
    _SECRETS = st.secrets if hasattr(st, "secrets") else {}
except Exception:  # pragma: no cover - allows core/ to be imported outside Streamlit
    _SECRETS = {}


def _get(key: str, default: str | None = None) -> str | None:
    """Look up a config value: env var first, then st.secrets, then default."""
    val = os.environ.get(key)
    if val:
        return val
    try:
        if key in _SECRETS:
            return _SECRETS[key]
    except Exception:
        pass
    return default


BASE_DIR = Path(__file__).resolve().parent
DATA_DIR = BASE_DIR / "data"
CHROMA_DIR = str(DATA_DIR / "chroma_store")
CACHE_DIR = DATA_DIR / "cache"
DATA_DIR.mkdir(parents=True, exist_ok=True)
CACHE_DIR.mkdir(parents=True, exist_ok=True)


@dataclass
class Settings:
    # --- LLM providers ---
    gemini_api_key: str | None = field(default_factory=lambda: _get("GEMINI_API_KEY"))
    groq_api_key: str | None = field(default_factory=lambda: _get("GROQ_API_KEY"))
    llm_provider: str = field(default_factory=lambda: _get("LLM_PROVIDER", "auto"))  # "gemini" | "groq" | "auto"

    gemini_model: str = "gemini-2.0-flash"
    groq_model: str = "openai/gpt-oss-120b"

    # --- Embeddings (local, no API needed) ---
    embed_model_name: str = "all-MiniLM-L6-v2"

    # --- Chunking ---
    chunk_size: int = 900
    chunk_overlap: int = 150
    top_k: int = 5

    # --- External threat intel APIs ---
    nvd_api_key: str | None = field(default_factory=lambda: _get("NVD_API_KEY"))  # optional, raises rate limit
    nvd_base_url: str = "https://services.nvd.nist.gov/rest/json/cves/2.0"
    kev_url: str = "https://www.cisa.gov/sites/default/files/feeds/known_exploited_vulnerabilities.json"

    # --- Cache TTLs (seconds) ---
    cve_cache_ttl: int = 6 * 3600
    kev_cache_ttl: int = 6 * 3600

    # --- Vector store ---
    chroma_dir: str = CHROMA_DIR
    cve_collection: str = "cve_intel"
    report_collection: str = "uploaded_reports"

    def has_llm(self) -> bool:
        return bool(self.gemini_api_key or self.groq_api_key)


settings = Settings()


# ============================================================================
# SECTION: core/llm.py
# ============================================================================

import logging
from dataclasses import dataclass


logger = logging.getLogger("sentry")


@dataclass
class LLMResponse:
    text: str
    provider: str
    model: str


class LLMUnavailableError(RuntimeError):
    pass


class LLMClient:
    """Thin wrapper that tries Gemini first, then Groq, based on settings.llm_provider."""

    def __init__(self):
        self._gemini_model = None
        self._groq_client = None
        self._init_gemini()
        self._init_groq()

    def _init_gemini(self):
        if not settings.gemini_api_key:
            return
        try:
            import google.generativeai as genai

            genai.configure(api_key=settings.gemini_api_key)
            self._gemini_model = genai.GenerativeModel(settings.gemini_model)
        except Exception as exc:  # pragma: no cover
            logger.warning("Gemini init failed: %s", exc)
            self._gemini_model = None

    def _init_groq(self):
        if not settings.groq_api_key:
            return
        try:
            from groq import Groq

            self._groq_client = Groq(api_key=settings.groq_api_key)
        except Exception as exc:  # pragma: no cover
            logger.warning("Groq init failed: %s", exc)
            self._groq_client = None

    @property
    def available(self) -> bool:
        return bool(self._gemini_model or self._groq_client)

    def generate(self, prompt: str, temperature: float = 0.2, max_tokens: int = 1500) -> LLMResponse:
        provider_order = []
        if settings.llm_provider == "gemini":
            provider_order = ["gemini"]
        elif settings.llm_provider == "groq":
            provider_order = ["groq"]
        else:  # auto
            provider_order = ["gemini", "groq"]

        last_error: Exception | None = None
        for provider in provider_order:
            try:
                if provider == "gemini" and self._gemini_model:
                    resp = self._gemini_model.generate_content(
                        prompt,
                        generation_config={"temperature": temperature, "max_output_tokens": max_tokens},
                    )
                    text = (resp.text or "").strip()
                    if text:
                        return LLMResponse(text=text, provider="gemini", model=settings.gemini_model)
                if provider == "groq" and self._groq_client:
                    completion = self._groq_client.chat.completions.create(
                        model=settings.groq_model,
                        messages=[{"role": "user", "content": prompt}],
                        temperature=temperature,
                        max_tokens=max_tokens,
                    )
                    text = completion.choices[0].message.content.strip()
                    if text:
                        return LLMResponse(text=text, provider="groq", model=settings.groq_model)
            except Exception as exc:
                logger.warning("Provider %s failed: %s", provider, exc)
                last_error = exc
                continue

        raise LLMUnavailableError(
            f"No LLM provider produced a response. Last error: {last_error}"
            if last_error
            else "No LLM provider is configured. Add GEMINI_API_KEY or GROQ_API_KEY."
        )


_client: LLMClient | None = None


def get_llm_client() -> LLMClient:
    global _client
    if _client is None:
        _client = LLMClient()
    return _client


def reset_llm_client():
    """Call after the user updates API keys in Settings so new keys take effect."""
    global _client
    _client = None


# ============================================================================
# SECTION: core/vectorstore.py
# ============================================================================


import time
import uuid
from typing import Any

import chromadb
import numpy as np
from sentence_transformers import SentenceTransformer


logger = logging.getLogger("sentry")

_embedder: SentenceTransformer | None = None
_chroma_client = None


def get_embedder() -> SentenceTransformer:
    global _embedder
    if _embedder is None:
        _embedder = SentenceTransformer(settings.embed_model_name)
    return _embedder


def embed_texts(texts: list[str]) -> list[list[float]]:
    vectors = get_embedder().encode(texts, normalize_embeddings=True)
    return np.asarray(vectors, dtype="float32").tolist()


def get_chroma_client():
    global _chroma_client
    if _chroma_client is None:
        _chroma_client = chromadb.PersistentClient(path=settings.chroma_dir)
    return _chroma_client


def get_collection(name: str):
    client = get_chroma_client()
    return client.get_or_create_collection(name=name, metadata={"hnsw:space": "cosine"})


class VectorStore:
    """High-level operations used by the rest of the app."""

    def __init__(self):
        self.cve_col = get_collection(settings.cve_collection)
        self.report_col = get_collection(settings.report_collection)

    # ------------------------------------------------------------------ CVE cache
    def cve_cached(self, cve_id: str) -> dict[str, Any] | None:
        try:
            result = self.cve_col.get(ids=[cve_id.upper()], include=["documents", "metadatas"])
        except Exception:
            return None
        if not result or not result.get("ids"):
            return None
        meta = result["metadatas"][0] or {}
        if time.time() - float(meta.get("cached_at", 0)) > settings.cve_cache_ttl:
            return None  # stale
        return {"text": result["documents"][0], "metadata": meta}

    def upsert_cve(self, cve_id: str, text: str, metadata: dict[str, Any]):
        metadata = {**metadata, "cached_at": time.time(), "type": "cve"}
        vector = embed_texts([text])[0]
        self.cve_col.upsert(ids=[cve_id.upper()], documents=[text], metadatas=[metadata], embeddings=[vector])

    def all_cached_cves(self) -> list[dict[str, Any]]:
        try:
            result = self.cve_col.get(include=["documents", "metadatas"])
        except Exception:
            return []
        out = []
        for cid, doc, meta in zip(result.get("ids", []), result.get("documents", []), result.get("metadatas", [])):
            out.append({"cve_id": cid, "text": doc, **(meta or {})})
        return out

    # ------------------------------------------------------------- uploaded reports
    def add_report_chunks(self, chunks: list[str], source_label: str, extra_meta: dict[str, Any] | None = None) -> int:
        if not chunks:
            return 0
        ids = [f"{source_label}-{uuid.uuid4().hex[:8]}-{i}" for i in range(len(chunks))]
        metadatas = [{"source": source_label, "chunk_index": i, "uploaded_at": time.time(), **(extra_meta or {})} for i in range(len(chunks))]
        vectors = embed_texts(chunks)
        self.report_col.upsert(ids=ids, documents=chunks, metadatas=metadatas, embeddings=vectors)
        return len(chunks)

    def report_sources(self) -> list[str]:
        try:
            result = self.report_col.get(include=["metadatas"])
        except Exception:
            return []
        return sorted({m.get("source", "unknown") for m in result.get("metadatas", [])})

    def query(self, collection, query_text: str, k: int = 5) -> list[dict[str, Any]]:
        if collection.count() == 0:
            return []
        vector = embed_texts([query_text])[0]
        result = collection.query(query_embeddings=[vector], n_results=min(k, collection.count()), include=["documents", "metadatas", "distances"])
        out = []
        docs = result.get("documents", [[]])[0]
        metas = result.get("metadatas", [[]])[0]
        dists = result.get("distances", [[]])[0]
        for doc, meta, dist in zip(docs, metas, dists):
            similarity = max(0.0, 1.0 - dist)  # cosine distance -> similarity
            out.append({"text": doc, "metadata": meta, "similarity": similarity})
        return out

    def query_reports(self, query_text: str, k: int = 5) -> list[dict[str, Any]]:
        return self.query(self.report_col, query_text, k)

    def query_cves(self, query_text: str, k: int = 5) -> list[dict[str, Any]]:
        return self.query(self.cve_col, query_text, k)


_store: VectorStore | None = None


def get_store() -> VectorStore:
    global _store
    if _store is None:
        _store = VectorStore()
    return _store


# ============================================================================
# SECTION: core/nvd_client.py
# ============================================================================

import json

import re


from pathlib import Path

import requests


logger = logging.getLogger("sentry")

CVE_PATTERN = re.compile(r"CVE-\d{4}-\d{4,7}", re.IGNORECASE)
_KEV_CACHE_FILE = Path(settings.chroma_dir).parent / "cache" / "kev_catalog.json"


@dataclass
class CVERecord:
    cve_id: str
    description: str = ""
    published: str = ""
    last_modified: str = ""
    cvss_score: float | None = None
    cvss_severity: str = "UNKNOWN"
    cvss_vector: str = ""
    cwe_ids: list[str] = field(default_factory=list)
    affected_products: list[str] = field(default_factory=list)
    references: list[str] = field(default_factory=list)
    vendor_advisory: str | None = None
    exploited_kev: bool = False
    kev_date_added: str | None = None
    kev_due_date: str | None = None
    patch_available: bool | None = None

    def to_text(self) -> str:
        """Flattened text used for embedding / LLM context."""
        lines = [
            f"CVE ID: {self.cve_id}",
            f"Published: {self.published}",
            f"Last Modified: {self.last_modified}",
            f"CVSS Score: {self.cvss_score if self.cvss_score is not None else 'N/A'} ({self.cvss_severity})",
            f"CVSS Vector: {self.cvss_vector or 'N/A'}",
            f"CWE: {', '.join(self.cwe_ids) if self.cwe_ids else 'N/A'}",
            f"Affected Products: {', '.join(self.affected_products[:15]) if self.affected_products else 'N/A'}",
            f"Actively Exploited (CISA KEV): {'YES' if self.exploited_kev else 'No known evidence'}",
            f"Description: {self.description or 'N/A'}",
            f"References: {', '.join(self.references[:8]) if self.references else 'N/A'}",
        ]
        if self.vendor_advisory:
            lines.append(f"Vendor Advisory: {self.vendor_advisory}")
        return "\n".join(lines)


def extract_cve_ids(text: str) -> list[str]:
    return sorted(set(m.upper() for m in CVE_PATTERN.findall(text or "")))


def _nvd_headers() -> dict[str, str]:
    headers = {"User-Agent": "SENTRY-Threat-Intel-Assistant/1.0"}
    if settings.nvd_api_key:
        headers["apiKey"] = settings.nvd_api_key
    return headers


def fetch_cve(cve_id: str, timeout: int = 15) -> CVERecord | None:
    """Fetch a single CVE's full detail from the NVD API. Returns None if not found."""
    cve_id = cve_id.strip().upper()
    if not CVE_PATTERN.fullmatch(cve_id):
        return None
    try:
        resp = requests.get(
            settings.nvd_base_url, params={"cveId": cve_id}, headers=_nvd_headers(), timeout=timeout
        )
        resp.raise_for_status()
        data = resp.json()
    except Exception as exc:
        logger.warning("NVD lookup failed for %s: %s", cve_id, exc)
        raise

    vulns = data.get("vulnerabilities", [])
    if not vulns:
        return None
    cve = vulns[0]["cve"]

    description = next((d["value"] for d in cve.get("descriptions", []) if d.get("lang") == "en"), "")

    metrics = cve.get("metrics", {})
    cvss_score, cvss_severity, cvss_vector = None, "UNKNOWN", ""
    for key in ("cvssMetricV31", "cvssMetricV30", "cvssMetricV2"):
        if key in metrics and metrics[key]:
            m = metrics[key][0]
            cvss_score = m["cvssData"].get("baseScore")
            cvss_severity = m.get("baseSeverity") or m["cvssData"].get("baseSeverity", "UNKNOWN")
            cvss_vector = m["cvssData"].get("vectorString", "")
            break

    cwe_ids = []
    for weakness in cve.get("weaknesses", []):
        for d in weakness.get("description", []):
            if d.get("value", "").startswith("CWE-"):
                cwe_ids.append(d["value"])

    affected_products = []
    for config in cve.get("configurations", []):
        for node in config.get("nodes", []):
            for cpe_match in node.get("cpeMatch", []):
                if cpe_match.get("vulnerable"):
                    criteria = cpe_match.get("criteria", "")
                    parts = criteria.split(":")
                    if len(parts) >= 6:
                        vendor, product = parts[3], parts[4]
                        affected_products.append(f"{vendor}:{product}")
    affected_products = sorted(set(affected_products))

    references = [r.get("url", "") for r in cve.get("references", []) if r.get("url")]
    vendor_advisory = next(
        (r for r in cve.get("references", []) if "Vendor Advisory" in r.get("tags", [])), None
    )
    vendor_advisory_url = vendor_advisory.get("url") if vendor_advisory else None

    record = CVERecord(
        cve_id=cve_id,
        description=description,
        published=cve.get("published", ""),
        last_modified=cve.get("lastModified", ""),
        cvss_score=cvss_score,
        cvss_severity=cvss_severity,
        cvss_vector=cvss_vector,
        cwe_ids=cwe_ids,
        affected_products=affected_products,
        references=references[:10],
        vendor_advisory=vendor_advisory_url,
    )

    kev_entry = get_kev_entry(cve_id)
    if kev_entry:
        record.exploited_kev = True
        record.kev_date_added = kev_entry.get("dateAdded")
        record.kev_due_date = kev_entry.get("dueDate")
        record.patch_available = True  # KEV entries always have a required action / patch path
    return record


# ---------------------------------------------------------------------------
# CISA KEV catalog
# ---------------------------------------------------------------------------
def get_kev_catalog(force_refresh: bool = False) -> list[dict[str, Any]]:
    """Return the full CISA KEV catalog, using a disk cache to avoid refetching the
    multi-MB feed on every call."""
    _KEV_CACHE_FILE.parent.mkdir(parents=True, exist_ok=True)
    if not force_refresh and _KEV_CACHE_FILE.exists():
        age = time.time() - _KEV_CACHE_FILE.stat().st_mtime
        if age < settings.kev_cache_ttl:
            try:
                return json.loads(_KEV_CACHE_FILE.read_text()).get("vulnerabilities", [])
            except Exception:
                pass
    try:
        resp = requests.get(settings.kev_url, timeout=20)
        resp.raise_for_status()
        data = resp.json()
        _KEV_CACHE_FILE.write_text(json.dumps(data))
        return data.get("vulnerabilities", [])
    except Exception as exc:
        logger.warning("CISA KEV fetch failed: %s", exc)
        if _KEV_CACHE_FILE.exists():
            try:
                return json.loads(_KEV_CACHE_FILE.read_text()).get("vulnerabilities", [])
            except Exception:
                pass
        return []


def get_kev_entry(cve_id: str) -> dict[str, Any] | None:
    catalog = get_kev_catalog()
    cve_id = cve_id.upper()
    for entry in catalog:
        if entry.get("cveID", "").upper() == cve_id:
            return entry
    return None


# ============================================================================
# SECTION: core/risk_engine.py
# ============================================================================

from dataclasses import dataclass
from datetime import datetime, timezone


# Weights sum to 100
WEIGHTS = {
    "cvss": 35,
    "active_exploitation": 25,
    "public_exploit": 15,
    "vendor_severity": 10,
    "patch_availability": -10,  # patched = lower risk (subtracted)
    "exploit_maturity": 10,
    "vulnerability_age": 5,
}


@dataclass
class RiskAssessment:
    score: int
    priority: str  # Critical | High | Medium | Low
    rationale: list[str]


def _cvss_component(cve: CVERecord) -> tuple[float, str]:
    if cve.cvss_score is None:
        return 0.0, "No CVSS score available — treated as unknown severity (0 pts)."
    pct = cve.cvss_score / 10.0
    pts = pct * WEIGHTS["cvss"]
    return pts, f"CVSS base score {cve.cvss_score} ({cve.cvss_severity}) contributes {pts:.1f}/{WEIGHTS['cvss']} pts."


def _exploitation_component(cve: CVERecord) -> tuple[float, str]:
    if cve.exploited_kev:
        pts = WEIGHTS["active_exploitation"]
        return pts, f"Listed in CISA's Known Exploited Vulnerabilities catalog — full {pts} pts (active exploitation confirmed)."
    return 0.0, "Not present in the CISA KEV catalog — no confirmed active exploitation (0 pts)."


def _public_exploit_component(cve: CVERecord) -> tuple[float, str]:
    exploit_refs = [r for r in cve.references if any(tag in r.lower() for tag in ["exploit-db", "exploit_db", "metasploit", "github.com", "packetstorm"])]
    if exploit_refs:
        pts = WEIGHTS["public_exploit"]
        return pts, f"{len(exploit_refs)} reference(s) point to likely public exploit code/PoC — {pts} pts."
    return 0.0, "No public exploit/PoC references detected in NVD data (0 pts)."


def _vendor_severity_component(cve: CVERecord) -> tuple[float, str]:
    if cve.vendor_advisory:
        pts = WEIGHTS["vendor_severity"]
        return pts, f"Vendor advisory published — {pts} pts (vendor has formally acknowledged the issue)."
    return WEIGHTS["vendor_severity"] * 0.3, "No vendor advisory found in NVD references — partial credit only."


def _patch_component(cve: CVERecord) -> tuple[float, str]:
    if cve.patch_available:
        pts = WEIGHTS["patch_availability"]
        return pts, f"Patch/remediation appears available — risk reduced by {abs(pts)} pts."
    return 0.0, "No confirmed patch availability signal — risk not reduced."


def _exploit_maturity_component(cve: CVERecord) -> tuple[float, str]:
    if cve.exploited_kev and cve.kev_due_date:
        pts = WEIGHTS["exploit_maturity"]
        return pts, f"CISA has set a remediation due date ({cve.kev_due_date}), indicating mature, weaponized exploitation — {pts} pts."
    if cve.exploited_kev:
        pts = WEIGHTS["exploit_maturity"] * 0.7
        return pts, f"Confirmed exploited but no due date on record — {pts:.1f} pts."
    return 0.0, "No indicators of exploit maturity beyond CVSS (0 pts)."


def _age_component(cve: CVERecord) -> tuple[float, str]:
    if not cve.published:
        return 0.0, "Publish date unavailable — age not scored."
    try:
        raw = cve.published.replace("Z", "+00:00")
        published = datetime.fromisoformat(raw)
        if published.tzinfo is None:
            published = published.replace(tzinfo=timezone.utc)
        age_days = (datetime.now(timezone.utc) - published).days
    except Exception:
        return 0.0, "Publish date unparseable — age not scored."
    if age_days > 365 and not cve.patch_available:
        pts = WEIGHTS["vulnerability_age"]
        return pts, f"Vulnerability is {age_days} days old and still unpatched — {pts} pts (long exposure window)."
    if age_days > 90:
        pts = WEIGHTS["vulnerability_age"] * 0.5
        return pts, f"Vulnerability is {age_days} days old — {pts:.1f} pts."
    return 0.0, f"Recently disclosed ({age_days} days old) — 0 pts."


def priority_from_score(score: int) -> str:
    if score >= 80:
        return "Critical"
    if score >= 60:
        return "High"
    if score >= 35:
        return "Medium"
    return "Low"


def assess_risk(cve: CVERecord) -> RiskAssessment:
    components = [
        _cvss_component(cve),
        _exploitation_component(cve),
        _public_exploit_component(cve),
        _vendor_severity_component(cve),
        _patch_component(cve),
        _exploit_maturity_component(cve),
        _age_component(cve),
    ]
    raw_score = sum(pts for pts, _ in components)
    score = max(0, min(100, round(raw_score)))
    rationale = [msg for _, msg in components]
    return RiskAssessment(score=score, priority=priority_from_score(score), rationale=rationale)


# ============================================================================
# SECTION: core/confidence.py
# ============================================================================


@dataclass
class ConfidenceResult:
    score: int  # 0-100
    label: str  # High | Medium | Low
    num_sources: int
    source_list: list[str]


def compute_confidence(retrieved: list[dict], has_live_cve_data: bool = False) -> ConfidenceResult:
    if not retrieved and not has_live_cve_data:
        return ConfidenceResult(score=0, label="Low", num_sources=0, source_list=[])

    similarities = [r.get("similarity", 0.0) for r in retrieved] or [0.0]
    mean_sim = sum(similarities) / len(similarities)

    sources = set()
    for r in retrieved:
        src = r.get("metadata", {}).get("source") or r.get("metadata", {}).get("cve_id") or "unknown"
        sources.add(src)
    if has_live_cve_data:
        sources.add("NVD/CISA")

    volume_factor = min(1.0, len(retrieved) / 4)  # 4+ chunks = full credit
    diversity_bonus = min(0.15, 0.05 * len(sources))

    raw = (mean_sim * 0.7) + (volume_factor * 0.15) + diversity_bonus
    score = max(0, min(100, round(raw * 100)))

    if score >= 75:
        label = "High"
    elif score >= 45:
        label = "Medium"
    else:
        label = "Low"

    return ConfidenceResult(score=score, label=label, num_sources=len(sources), source_list=sorted(sources))


# ============================================================================
# SECTION: core/ingestion.py
# ============================================================================

import io


import re

import pandas as pd


logger = logging.getLogger("sentry")

SUPPORTED_EXTENSIONS = {"pdf", "docx", "txt", "csv", "xlsx", "json"}


def extract_text(file_bytes: bytes, filename: str) -> str:
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    if ext not in SUPPORTED_EXTENSIONS:
        raise ValueError(f"Unsupported file type: .{ext}. Supported: {sorted(SUPPORTED_EXTENSIONS)}")

    if ext == "pdf":
        return _extract_pdf(file_bytes)
    if ext == "docx":
        return _extract_docx(file_bytes)
    if ext == "txt":
        return file_bytes.decode("utf-8", errors="ignore")
    if ext == "csv":
        return _extract_tabular(file_bytes, fmt="csv")
    if ext == "xlsx":
        return _extract_tabular(file_bytes, fmt="xlsx")
    if ext == "json":
        return _extract_json(file_bytes)
    raise ValueError(f"Unhandled extension: {ext}")  # pragma: no cover


def _extract_pdf(file_bytes: bytes) -> str:
    import fitz  # PyMuPDF

    doc = fitz.open(stream=file_bytes, filetype="pdf")
    return "\n".join(page.get_text() for page in doc)


def _extract_docx(file_bytes: bytes) -> str:
    import docx  # python-docx

    document = docx.Document(io.BytesIO(file_bytes))
    parts = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def _extract_tabular(file_bytes: bytes, fmt: str) -> str:
    if fmt == "csv":
        df = pd.read_csv(io.BytesIO(file_bytes))
    else:
        df = pd.read_excel(io.BytesIO(file_bytes))
    lines = [" | ".join(str(c) for c in df.columns)]
    for _, row in df.iterrows():
        lines.append(" | ".join(str(v) for v in row.values))
    return "\n".join(lines)


def _extract_json(file_bytes: bytes) -> str:
    data = json.loads(file_bytes.decode("utf-8", errors="ignore"))
    return json.dumps(data, indent=2)


def chunk_text(text: str, size: int | None = None, overlap: int | None = None) -> list[str]:
    size = size or settings.chunk_size
    overlap = overlap or settings.chunk_overlap
    text = re.sub(r"\s+", " ", text).strip()
    chunks = []
    start = 0
    while start < len(text):
        end = start + size
        chunks.append(text[start:end])
        start += size - overlap
    return [c for c in chunks if len(c.strip()) > 30]


# ============================================================================
# SECTION: core/prompts.py
# ============================================================================

INTENT_PROMPT = """Classify the analyst's request into exactly one category. Reply with ONLY the
category word, nothing else.

Categories:
- cve_query: asks about one or more specific CVE IDs, or general CVE lookup
- comparison: asks to compare two or more CVEs
- document_query: asks about an uploaded report (summarize, extract CVEs, recommendations, executive summary)
- filter_query: asks to filter/list vulnerabilities by severity, vendor, product, or exploitation status
- general: anything else (definitions, general security questions)

Request: "{query}"

Category:"""


TECHNICAL_ANSWER_PROMPT = """You are SENTRY, a senior cyber security threat-intelligence analyst assistant.
Answer the analyst's question using ONLY the context provided below. If the context does not
contain the answer, say so explicitly rather than guessing or inventing information.

Respond in TECHNICAL mode:
- Use precise security terminology (attack vector, CWE class, exploitation prerequisites, etc.)
- Include affected software/versions, mitigation steps, and references where available
- Cite the source label (e.g. "[Source: CVE-2024-3400]" or "[Source: uploaded_report.pdf]") for each claim

Context:
{context}

Question: {query}

Technical Answer:"""


BEGINNER_ANSWER_PROMPT = """You are SENTRY, a friendly security assistant explaining things to a
non-technical reader. Answer the question using ONLY the context provided below. If the context
does not contain the answer, say so clearly rather than guessing.

Respond in BEGINNER mode:
- Avoid jargon (no "RCE", "CWE", "attack surface" without plain-language explanation)
- Use short sentences and simple analogies
- Explain what could go wrong in practical terms and what the reader should do about it
- Still mention the source of your information in plain language (e.g. "according to the vendor's advisory")

Context:
{context}

Question: {query}

Simple Explanation:"""


SUMMARY_PROMPT = """You are SENTRY, a cyber security analyst assistant. Using ONLY the report
context below, produce an executive summary suitable for a CISO. Include: overall risk posture,
number and severity breakdown of findings, top 3 priorities to patch first, and key recommendations.
If the context is insufficient for any section, state that explicitly.

Report context:
{context}

Executive Summary:"""


def build_answer_prompt(query: str, context: str, mode: str) -> str:
    template = TECHNICAL_ANSWER_PROMPT if mode == "technical" else BEGINNER_ANSWER_PROMPT
    return template.format(context=context or "No relevant context was retrieved.", query=query)


# ============================================================================
# SECTION: core/graph.py
# ============================================================================


from typing import Any, TypedDict

from langgraph.graph import END, StateGraph


logger = logging.getLogger("sentry")


class SentryState(TypedDict, total=False):
    query: str
    mode: str  # "technical" | "beginner"
    intent: str
    cve_ids: list[str]
    cve_records: dict[str, CVERecord]
    retrieved: list[dict[str, Any]]
    context: str
    answer: str
    risk: dict[str, RiskAssessment]
    confidence: Any
    errors: list[str]


def node_detect_intent(state: SentryState) -> SentryState:
    query = state["query"]
    cve_ids = extract_cve_ids(query)
    state["cve_ids"] = cve_ids

    if cve_ids and len(cve_ids) >= 2 and ("compar" in query.lower() or " vs " in query.lower()):
        state["intent"] = "comparison"
    elif cve_ids:
        state["intent"] = "cve_query"
    elif any(kw in query.lower() for kw in ["this report", "uploaded", "summarize", "summary", "executive summary", "extract all cve", "recommendations"]):
        state["intent"] = "document_query"
    elif any(kw in query.lower() for kw in ["only critical", "which vulnerabilities", "affect windows", "affect apache", "patch first", "filter"]):
        state["intent"] = "filter_query"
    else:
        state["intent"] = "general"
    return state


def node_retrieve_live(state: SentryState) -> SentryState:
    store = get_store()
    records: dict[str, CVERecord] = {}
    errors = state.get("errors", [])
    for cve_id in state.get("cve_ids", []):
        cached = store.cve_cached(cve_id)
        if cached:
            # Reconstruct a minimal record from cached metadata + text for downstream use.
            meta = cached["metadata"]
            record = CVERecord(
                cve_id=cve_id,
                description=cached["text"],
                cvss_score=meta.get("cvss_score"),
                cvss_severity=meta.get("cvss_severity", "UNKNOWN"),
                exploited_kev=bool(meta.get("exploited_kev", False)),
                patch_available=meta.get("patch_available"),
            )
            records[cve_id] = record
            continue
        try:
            record = fetch_cve(cve_id)
        except Exception as exc:
            errors.append(f"NVD lookup failed for {cve_id}: {exc}")
            continue
        if record is None:
            errors.append(f"No NVD data found for {cve_id}.")
            continue
        records[cve_id] = record
        store.upsert_cve(
            cve_id,
            record.to_text(),
            {
                "cvss_score": record.cvss_score,
                "cvss_severity": record.cvss_severity,
                "exploited_kev": record.exploited_kev,
                "patch_available": record.patch_available,
            },
        )
    state["cve_records"] = records
    state["errors"] = errors
    return state


def node_retrieve_reports(state: SentryState) -> SentryState:
    store = get_store()
    results = store.query_reports(state["query"], k=5)
    state["retrieved"] = state.get("retrieved", []) + results
    return state


def node_retrieve_cve_cache(state: SentryState) -> SentryState:
    store = get_store()
    results = store.query_cves(state["query"], k=3)
    state["retrieved"] = state.get("retrieved", []) + results
    return state


def node_merge_context(state: SentryState) -> SentryState:
    parts = []
    for cve_id, record in state.get("cve_records", {}).items():
        parts.append(f"[Source: {cve_id} (NVD/CISA KEV)]\n{record.to_text()}")
    for r in state.get("retrieved", []):
        src = r["metadata"].get("source") or r["metadata"].get("cve_id", "unknown")
        parts.append(f"[Source: {src}]\n{r['text']}")
    state["context"] = "\n\n---\n\n".join(parts)
    return state


def node_generate_answer(state: SentryState) -> SentryState:
    client = get_llm_client()
    prompt = prompts.build_answer_prompt(state["query"], state.get("context", ""), state.get("mode", "technical"))
    try:
        response = client.generate(prompt)
        state["answer"] = response.text
    except Exception as exc:
        state["answer"] = (
            "I couldn't generate an answer because no LLM provider is available or the call failed "
            f"({exc}). Please check your API key in Settings."
        )
    return state


def node_compute_risk(state: SentryState) -> SentryState:
    risk: dict[str, RiskAssessment] = {}
    for cve_id, record in state.get("cve_records", {}).items():
        risk[cve_id] = assess_risk(record)
    state["risk"] = risk
    return state


def node_compute_confidence(state: SentryState) -> SentryState:
    has_live = bool(state.get("cve_records"))
    state["confidence"] = compute_confidence(state.get("retrieved", []), has_live_cve_data=has_live)
    return state


def route_after_intent(state: SentryState) -> str:
    if state["intent"] in ("cve_query", "comparison", "filter_query") and state.get("cve_ids"):
        return "cve_path"
    return "doc_path"


def build_graph():
    graph = StateGraph(SentryState)
    graph.add_node("detect_intent", node_detect_intent)
    graph.add_node("retrieve_live", node_retrieve_live)
    graph.add_node("retrieve_cve_cache", node_retrieve_cve_cache)
    graph.add_node("retrieve_reports", node_retrieve_reports)
    graph.add_node("merge_context", node_merge_context)
    graph.add_node("generate_answer", node_generate_answer)
    graph.add_node("compute_risk", node_compute_risk)
    graph.add_node("compute_confidence", node_compute_confidence)

    graph.set_entry_point("detect_intent")
    graph.add_conditional_edges(
        "detect_intent", route_after_intent, {"cve_path": "retrieve_live", "doc_path": "retrieve_reports"}
    )
    graph.add_edge("retrieve_live", "retrieve_cve_cache")
    graph.add_edge("retrieve_cve_cache", "retrieve_reports")
    graph.add_edge("retrieve_reports", "merge_context")
    graph.add_edge("merge_context", "generate_answer")
    graph.add_edge("generate_answer", "compute_risk")
    graph.add_edge("compute_risk", "compute_confidence")
    graph.add_edge("compute_confidence", END)

    return graph.compile()


_compiled_graph = None


def get_graph():
    global _compiled_graph
    if _compiled_graph is None:
        _compiled_graph = build_graph()
    return _compiled_graph


def run_investigation(query: str, mode: str = "technical") -> SentryState:
    initial: SentryState = {"query": query, "mode": mode, "errors": []}
    graph = get_graph()
    result = graph.invoke(initial)
    return result


# ============================================================================
# SECTION: core/report_generator.py
# ============================================================================


@dataclass
class InvestigationReport:
    question: str
    answer: str
    cve_id: str | None = None
    severity: str = "Unknown"
    cvss_score: float | None = None
    risk_score: int | None = None
    priority: str = "N/A"
    risk_rationale: list[str] = field(default_factory=list)
    exploited_kev: bool = False
    affected_products: list[str] = field(default_factory=list)
    patch_available: bool | None = None
    mitigation: str = ""
    references: list[str] = field(default_factory=list)
    confidence_score: int = 0
    confidence_label: str = "Low"
    sources: list[str] = field(default_factory=list)
    generated_at: str = field(default_factory=lambda: datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC"))


_SEVERITY_COLORS = {
    "Critical": (0xE5, 0x1C, 0x2B),
    "High": (0xE6, 0x7E, 0x22),
    "Medium": (0xE6, 0xB4, 0x22),
    "Low": (0x2E, 0xA0, 0x4B),
    "N/A": (0x77, 0x77, 0x77),
    "Unknown": (0x77, 0x77, 0x77),
}


def _color_for(priority: str):
    return _SEVERITY_COLORS.get(priority, _SEVERITY_COLORS["Unknown"])


def generate_pdf(report: InvestigationReport) -> bytes:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import cm
    from reportlab.platypus import (
        Paragraph,
        SimpleDocTemplate,
        Spacer,
        Table,
        TableStyle,
    )

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, topMargin=2 * cm, bottomMargin=2 * cm)
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("SentryTitle", parent=styles["Title"], textColor=colors.HexColor("#0D131C"))
    h2 = ParagraphStyle("SentryH2", parent=styles["Heading2"], textColor=colors.HexColor("#0F8F7C"))
    body = styles["BodyText"]

    r, g, b = _color_for(report.priority)
    sev_color = colors.Color(r / 255, g / 255, b / 255)

    story = [
        Paragraph("SENTRY — Threat Investigation Report", title_style),
        Spacer(1, 6),
        Paragraph(f"Generated: {report.generated_at}", body),
        Spacer(1, 14),
        Paragraph("Investigation Summary", h2),
        Paragraph(f"<b>Question:</b> {report.question}", body),
        Spacer(1, 6),
        Paragraph(report.answer.replace(chr(10), "<br/>"), body),
        Spacer(1, 14),
    ]

    if report.cve_id:
        story.append(Paragraph("Vulnerability Details", h2))
        details = [
            ["CVE ID", report.cve_id],
            ["CVSS Score", str(report.cvss_score) if report.cvss_score is not None else "N/A"],
            ["Severity", report.severity],
            ["Risk Score", f"{report.risk_score}/100" if report.risk_score is not None else "N/A"],
            ["Priority", report.priority],
            ["Actively Exploited (KEV)", "Yes" if report.exploited_kev else "No known evidence"],
            ["Patch Available", "Yes" if report.patch_available else "Unknown / No"],
            ["Affected Products", ", ".join(report.affected_products[:10]) or "N/A"],
        ]
        table = Table(details, colWidths=[5 * cm, 11 * cm])
        table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#131B27")),
                    ("TEXTCOLOR", (0, 0), (0, -1), colors.white),
                    ("BACKGROUND", (1, 4), (1, 4), sev_color),  # priority row highlight approx
                    ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#CCCCCC")),
                    ("FONTSIZE", (0, 0), (-1, -1), 9),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("TOPPADDING", (0, 0), (-1, -1), 4),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                ]
            )
        )
        story.append(table)
        story.append(Spacer(1, 14))

        if report.risk_rationale:
            story.append(Paragraph("Risk Score Rationale", h2))
            for line in report.risk_rationale:
                story.append(Paragraph(f"• {line}", body))
            story.append(Spacer(1, 14))

    if report.mitigation:
        story.append(Paragraph("Recommended Mitigation", h2))
        story.append(Paragraph(report.mitigation.replace(chr(10), "<br/>"), body))
        story.append(Spacer(1, 14))

    if report.references:
        story.append(Paragraph("References", h2))
        for ref in report.references[:10]:
            story.append(Paragraph(f"• {ref}", body))
        story.append(Spacer(1, 14))

    story.append(Paragraph("Confidence & Sources", h2))
    story.append(Paragraph(f"<b>Confidence:</b> {report.confidence_score}% ({report.confidence_label})", body))
    story.append(Paragraph(f"<b>Sources used:</b> {', '.join(report.sources) or 'N/A'}", body))

    doc.build(story)
    return buf.getvalue()


def generate_docx(report: InvestigationReport) -> bytes:
    import docx
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    document = docx.Document()

    title = document.add_heading("SENTRY — Threat Investigation Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    meta = document.add_paragraph(f"Generated: {report.generated_at}")
    meta.runs[0].italic = True

    document.add_heading("Investigation Summary", level=1)
    p = document.add_paragraph()
    p.add_run("Question: ").bold = True
    p.add_run(report.question)
    document.add_paragraph(report.answer)

    if report.cve_id:
        document.add_heading("Vulnerability Details", level=1)
        table = document.add_table(rows=0, cols=2)
        table.style = "Light Grid Accent 1"
        rows_data = [
            ("CVE ID", report.cve_id),
            ("CVSS Score", str(report.cvss_score) if report.cvss_score is not None else "N/A"),
            ("Severity", report.severity),
            ("Risk Score", f"{report.risk_score}/100" if report.risk_score is not None else "N/A"),
            ("Priority", report.priority),
            ("Actively Exploited (KEV)", "Yes" if report.exploited_kev else "No known evidence"),
            ("Patch Available", "Yes" if report.patch_available else "Unknown / No"),
            ("Affected Products", ", ".join(report.affected_products[:10]) or "N/A"),
        ]
        for label, value in rows_data:
            row = table.add_row().cells
            row[0].text = label
            row[1].text = value

        r, g, b = _color_for(report.priority)
        priority_para = document.add_paragraph()
        run = priority_para.add_run(f"Priority: {report.priority}")
        run.bold = True
        run.font.color.rgb = RGBColor(r, g, b)
        run.font.size = Pt(13)

        if report.risk_rationale:
            document.add_heading("Risk Score Rationale", level=2)
            for line in report.risk_rationale:
                document.add_paragraph(line, style="List Bullet")

    if report.mitigation:
        document.add_heading("Recommended Mitigation", level=1)
        document.add_paragraph(report.mitigation)

    if report.references:
        document.add_heading("References", level=1)
        for ref in report.references[:10]:
            document.add_paragraph(ref, style="List Bullet")

    document.add_heading("Confidence & Sources", level=1)
    document.add_paragraph(f"Confidence: {report.confidence_score}% ({report.confidence_label})")
    document.add_paragraph(f"Sources used: {', '.join(report.sources) or 'N/A'}")

    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


# ============================================================================
# SECTION: core/storage.py
# ============================================================================


from typing import Any


HISTORY_FILE = DATA_DIR / "investigation_history.jsonl"
REPORTS_DIR = DATA_DIR / "reports"
REPORTS_DIR.mkdir(parents=True, exist_ok=True)


def log_investigation(record: dict[str, Any]) -> str:
    record = {"id": uuid.uuid4().hex[:10], "timestamp": time.time(), **record}
    with HISTORY_FILE.open("a") as f:
        f.write(json.dumps(record, default=str) + "\n")
    return record["id"]


def load_history(limit: int = 200) -> list[dict[str, Any]]:
    if not HISTORY_FILE.exists():
        return []
    lines = HISTORY_FILE.read_text().splitlines()
    records = [json.loads(line) for line in lines if line.strip()]
    return list(reversed(records))[:limit]


def save_report_file(filename: str, content: bytes) -> Path:
    path = REPORTS_DIR / filename
    path.write_bytes(content)
    return path


def list_report_files() -> list[Path]:
    return sorted(REPORTS_DIR.glob("*"), key=lambda p: p.stat().st_mtime, reverse=True)



# ============================================================================
# SECTION: Streamlit UI
# ============================================================================
import time

import pandas as pd
import plotly.express as px
import streamlit as st


st.set_page_config(page_title="SENTRY // Threat Intel Console", page_icon="🛡️", layout="wide")

# ---------------------------------------------------------------------------
# Styling (SOC console theme)
# ---------------------------------------------------------------------------
st.markdown(
    """
<style>
@import url('https://fonts.googleapis.com/css2?family=Space+Grotesk:wght@500;600;700&family=JetBrains+Mono:wght@400;500;600;700&family=Inter:wght@400;500;600;700&display=swap');
:root{
  --bg:#05070B; --panel:#0D131C; --panel-2:#131B27; --border:#1F2B3B; --text:#E7EEF5; --muted:#7688A0;
  --accent:#2FE6C7; --accent-dim:#0F8F7C; --accent2:#7B61FF; --danger:#FF4D5E; --warn:#FFB627; --low:#4ADE80;
}
[data-testid="stAppViewContainer"]{
  background: radial-gradient(160% 120% at 15% -10%, rgba(123,97,255,0.07), transparent 55%),
    radial-gradient(120% 100% at 90% 0%, rgba(47,230,199,0.06), transparent 50%), var(--bg);
}
[data-testid="stHeader"]{ background: transparent; }
html, body, [class*="css"]{ font-family:'Inter', sans-serif; color: var(--text); }
h1,h2,h3,h4,h5{ font-family:'Space Grotesk', sans-serif; }
.sentry-hero{ border:1px solid var(--border); border-radius:18px; background: linear-gradient(180deg, #0B1119 0%, #070B11 100%); padding:26px 32px; margin-bottom:22px; }
.sentry-title{ font-family:'Space Grotesk', sans-serif; font-size:1.8rem; font-weight:700; margin:0; color:#F3F8FC; }
.sentry-sub{ font-size:0.92rem; color:var(--muted); margin-top:4px; }
.console-label{ font-family:'JetBrains Mono', monospace; font-size:0.74rem; letter-spacing:0.14em; color:var(--accent); text-transform:uppercase; margin:14px 0 8px 0; }
[data-testid="stSidebar"]{ background: var(--panel); border-right:1px solid var(--border); }
.stButton button, .stFormSubmitButton button, .stDownloadButton button{
  background: linear-gradient(145deg, var(--accent-dim), var(--accent)) !important; color:#03110D !important;
  font-weight:700 !important; border:none !important; border-radius:8px !important; font-family:'JetBrains Mono', monospace !important;
}
[data-testid="stChatMessage"]{ background: var(--panel) !important; border:1px solid var(--border) !important; border-radius:12px !important; }
.metric-card{ background:var(--panel); border:1px solid var(--border); border-radius:14px; padding:18px 20px; text-align:center; }
.metric-value{ font-family:'Space Grotesk',sans-serif; font-size:2rem; font-weight:700; }
.metric-label{ font-family:'JetBrains Mono',monospace; font-size:0.72rem; letter-spacing:0.08em; color:var(--muted); text-transform:uppercase; }
.sev-chip{ display:inline-flex; align-items:center; gap:6px; padding:3px 10px; border-radius:999px; font-family:'JetBrains Mono', monospace; font-size:0.72rem; font-weight:700; }
.sev-critical{ background:rgba(255,77,94,0.12); color:var(--danger); border:1px solid rgba(255,77,94,0.35); }
.sev-high{ background:rgba(255,182,39,0.12); color:var(--warn); border:1px solid rgba(255,182,39,0.35); }
.sev-medium{ background:rgba(255,182,39,0.10); color:var(--warn); border:1px solid rgba(255,182,39,0.25); }
.sev-low{ background:rgba(74,222,128,0.12); color:var(--low); border:1px solid rgba(74,222,128,0.35); }
.sev-unknown, .sev-na{ background:rgba(124,139,158,0.12); color:var(--muted); border:1px solid rgba(124,139,158,0.35); }
.src-chip{ display:inline-block; font-family:'JetBrains Mono', monospace; font-size:0.72rem; color:var(--accent);
  background:rgba(47,230,199,0.08); border:1px solid rgba(47,230,199,0.25); padding:2px 8px; border-radius:6px; margin:2px 4px 2px 0; }
.risk-meter-track{ width:100%; height:10px; background:var(--panel-2); border-radius:999px; overflow:hidden; border:1px solid var(--border); }
.risk-meter-fill{ height:100%; border-radius:999px; }
</style>
""",
    unsafe_allow_html=True,
)

PRIORITY_CLASS = {"Critical": "sev-critical", "High": "sev-high", "Medium": "sev-medium", "Low": "sev-low"}


def sev_chip(label: str, extra: str = "") -> str:
    css = PRIORITY_CLASS.get(label, "sev-unknown")
    suffix = f" · {extra}" if extra else ""
    return f'<span class="sev-chip {css}">{label}{suffix}</span>'


def risk_meter(score: int) -> str:
    color = "var(--danger)" if score >= 80 else "var(--warn)" if score >= 60 else "#F2C744" if score >= 35 else "var(--low)"
    return (
        f'<div class="risk-meter-track"><div class="risk-meter-fill" '
        f'style="width:{score}%;background:{color};"></div></div>'
    )


# ---------------------------------------------------------------------------
# Session state
# ---------------------------------------------------------------------------
for key, default in {
    "messages": [],
    "explanation_mode": "technical",
    "last_investigation": None,
    "nav": "Dashboard",
}.items():
    if key not in st.session_state:
        st.session_state[key] = default

store = get_store()
llm = get_llm_client()

# ---------------------------------------------------------------------------
# Sidebar navigation
# ---------------------------------------------------------------------------
with st.sidebar:
    st.markdown(
        '<div style="display:flex;align-items:center;gap:10px;margin-bottom:6px;">'
        '<span style="font-size:1.6rem;">🛡️</span>'
        '<span style="font-family:\'Space Grotesk\',sans-serif;font-weight:700;font-size:1.2rem;">SENTRY</span>'
        "</div>",
        unsafe_allow_html=True,
    )
    st.caption("Threat Intelligence RAG Assistant")
    st.session_state.nav = st.radio(
        "Navigate",
        ["Dashboard", "Investigate (Chat)", "Upload Report", "Investigation History", "Reports", "Settings"],
        label_visibility="collapsed",
    )
    st.divider()
    st.markdown('<p class="console-label">⚙️ Explanation Mode</p>', unsafe_allow_html=True)
    st.session_state.explanation_mode = st.radio(
        "Mode", ["technical", "beginner"], horizontal=True, label_visibility="collapsed",
        format_func=lambda m: "🧪 Technical" if m == "technical" else "🧑\u200d🎓 Beginner",
    )
    st.divider()
    st.markdown('<p class="console-label">📊 Mission Stats</p>', unsafe_allow_html=True)
    report_sources = store.report_sources()
    cached_cves = store.all_cached_cves()
    st.markdown(
        f"""
        <div style="font-family:'JetBrains Mono',monospace;font-size:0.78rem;color:var(--muted);">
        <div>Reports indexed: <b style="color:var(--text)">{len(report_sources)}</b></div>
        <div>CVEs cached: <b style="color:var(--text)">{len(cached_cves)}</b></div>
        <div>LLM ready: <b style="color:var(--text)">{"Yes" if llm.available else "No"}</b></div>
        </div>
        """,
        unsafe_allow_html=True,
    )
    if not llm.available:
        st.warning("No LLM configured. Go to Settings to add a Gemini or Groq API key.")

# ---------------------------------------------------------------------------
# Hero header
# ---------------------------------------------------------------------------
st.markdown(
    '<div class="sentry-hero"><p class="sentry-title">🛡️ SENTRY — Threat Intel Console</p>'
    '<p class="sentry-sub">Investigate CVEs, analyze vulnerability reports, and get prioritized, '
    "source-backed answers — powered by live NVD/CISA intel + your uploaded reports.</p></div>",
    unsafe_allow_html=True,
)


# =============================================================================
# PAGE: Dashboard
# =============================================================================
def page_dashboard():
    st.markdown('<p class="console-label">📊 Security Dashboard</p>', unsafe_allow_html=True)

    cves = store.all_cached_cves()
    kev_catalog = get_kev_catalog()
    kev_ids = {e.get("cveID") for e in kev_catalog}

    counts = {"Critical": 0, "High": 0, "Medium": 0, "Low": 0}
    rows = []
    for c in cves:
        score = c.get("cvss_score")
        severity = priority_from_score(int((score or 0) / 10 * 100)) if score is not None else "Low"
        counts[severity] = counts.get(severity, 0) + 1
        rows.append(
            {
                "CVE": c["cve_id"],
                "CVSS": score,
                "Severity": severity,
                "Exploited (KEV)": "Yes" if c["cve_id"] in kev_ids else "No",
                "Cached": pd.to_datetime(c.get("cached_at", 0), unit="s", errors="coerce"),
            }
        )
    df = pd.DataFrame(rows)

    cols = st.columns(4)
    for col, (label, value) in zip(cols, counts.items()):
        css = PRIORITY_CLASS.get(label, "sev-unknown")
        color_var = {"Critical": "var(--danger)", "High": "var(--warn)", "Medium": "#F2C744", "Low": "var(--low)"}[label]
        col.markdown(
            f'<div class="metric-card"><div class="metric-value" style="color:{color_var}">{value}</div>'
            f'<div class="metric-label">{label}</div></div>',
            unsafe_allow_html=True,
        )

    st.markdown("<br/>", unsafe_allow_html=True)

    if df.empty:
        st.info("No CVEs investigated yet. Head to **Investigate (Chat)** and ask about a CVE to populate the dashboard.")
        return

    c1, c2 = st.columns(2)
    with c1:
        st.markdown('<p class="console-label">Severity Distribution</p>', unsafe_allow_html=True)
        sev_df = df["Severity"].value_counts().reset_index()
        sev_df.columns = ["Severity", "Count"]
        fig = px.pie(
            sev_df, names="Severity", values="Count", hole=0.55,
            color="Severity",
            color_discrete_map={"Critical": "#FF4D5E", "High": "#FFB627", "Medium": "#F2C744", "Low": "#4ADE80"},
        )
        fig.update_layout(paper_bgcolor="rgba(0,0,0,0)", plot_bgcolor="rgba(0,0,0,0)", font_color="#E7EEF5", legend_title_text="")
        st.plotly_chart(fig, use_container_width=True)
    with c2:
        st.markdown('<p class="console-label">Recently Added CVEs</p>', unsafe_allow_html=True)
        recent = df.sort_values("Cached", ascending=False).head(10)
        fig2 = px.bar(recent, x="CVE", y="CVSS", color="Severity",
                      color_discrete_map={"Critical": "#FF4D5E", "High": "#FFB627", "Medium": "#F2C744", "Low": "#4ADE80"})
        fig2.update_layout(paper_bgcolor="rgba(0,0,0,0)", plot_bgcolor="rgba(0,0,0,0)", font_color="#E7EEF5")
        st.plotly_chart(fig2, use_container_width=True)

    st.markdown('<p class="console-label">🔍 Filter & Browse</p>', unsafe_allow_html=True)
    fc1, fc2, fc3 = st.columns(3)
    sev_filter = fc1.multiselect("Severity", ["Critical", "High", "Medium", "Low"], default=[])
    exploited_filter = fc2.selectbox("Exploitation Status", ["All", "Actively Exploited", "Not Exploited"])
    search_filter = fc3.text_input("Search CVE ID")

    filtered = df.copy()
    if sev_filter:
        filtered = filtered[filtered["Severity"].isin(sev_filter)]
    if exploited_filter == "Actively Exploited":
        filtered = filtered[filtered["Exploited (KEV)"] == "Yes"]
    elif exploited_filter == "Not Exploited":
        filtered = filtered[filtered["Exploited (KEV)"] == "No"]
    if search_filter:
        filtered = filtered[filtered["CVE"].str.contains(search_filter.upper(), na=False)]

    st.dataframe(filtered.drop(columns=["Cached"]), use_container_width=True, hide_index=True)


# =============================================================================
# PAGE: Investigate (Chat)
# =============================================================================
def page_chat():
    st.markdown('<p class="console-label">💬 Investigate</p>', unsafe_allow_html=True)
    st.caption(
        "Ask about a CVE (e.g. *Explain CVE-2024-3400*), compare CVEs, query an uploaded report, "
        "or filter (*show only Critical vulnerabilities*)."
    )

    for msg in st.session_state.messages:
        with st.chat_message(msg["role"]):
            st.markdown(msg["content"], unsafe_allow_html=True)

    query = st.chat_input("Ask SENTRY a question...")
    if not query:
        return

    if not llm.available:
        st.error("No LLM provider configured. Add a Gemini or Groq API key in Settings.")
        return

    st.session_state.messages.append({"role": "user", "content": query})
    with st.chat_message("user"):
        st.markdown(query)

    with st.chat_message("assistant"):
        with st.spinner("SENTRY is investigating..."):
            result = run_investigation(query, mode=st.session_state.explanation_mode)

        answer = result.get("answer", "No answer generated.")
        st.markdown(answer)

        for err in result.get("errors", []):
            st.warning(err)

        confidence: ConfidenceResult = result["confidence"]
        conf_color = "var(--low)" if confidence.label == "High" else "var(--warn)" if confidence.label == "Medium" else "var(--danger)"
        st.markdown(
            f'<div style="margin-top:8px;"><b>Confidence:</b> '
            f'<span style="color:{conf_color};font-weight:700;">{confidence.score}% ({confidence.label})</span> '
            f"— based on {confidence.num_sources} source(s), {len(result.get('retrieved', []))} retrieved chunk(s).</div>",
            unsafe_allow_html=True,
        )
        if confidence.source_list:
            chips = "".join(f'<span class="src-chip">{s}</span>' for s in confidence.source_list)
            st.markdown(chips, unsafe_allow_html=True)

        risk_data = result.get("risk", {})
        report_payload = None
        if risk_data:
            st.markdown("---")
            for cve_id, assessment in risk_data.items():
                record = result["cve_records"].get(cve_id)
                st.markdown(f"**{cve_id}** — {sev_chip(assessment.priority, f'{assessment.score}/100')}", unsafe_allow_html=True)
                st.markdown(risk_meter(assessment.score), unsafe_allow_html=True)
                with st.expander(f"Why this score? ({cve_id})"):
                    for line in assessment.rationale:
                        st.markdown(f"- {line}")

                report_payload = InvestigationReport(
                    question=query,
                    answer=answer,
                    cve_id=cve_id,
                    severity=record.cvss_severity if record else "Unknown",
                    cvss_score=record.cvss_score if record else None,
                    risk_score=assessment.score,
                    priority=assessment.priority,
                    risk_rationale=assessment.rationale,
                    exploited_kev=record.exploited_kev if record else False,
                    affected_products=record.affected_products if record else [],
                    patch_available=record.patch_available if record else None,
                    mitigation="Apply the vendor patch referenced above. If unavailable, apply vendor-recommended "
                    "workarounds and increase monitoring for exploitation indicators.",
                    references=record.references if record else [],
                    confidence_score=confidence.score,
                    confidence_label=confidence.label,
                    sources=confidence.source_list,
                )
        else:
            report_payload = InvestigationReport(
                question=query,
                answer=answer,
                confidence_score=confidence.score,
                confidence_label=confidence.label,
                sources=confidence.source_list,
            )

        with st.expander("📎 Retrieved sources"):
            for r in result.get("retrieved", []):
                src = r["metadata"].get("source") or r["metadata"].get("cve_id", "unknown")
                st.markdown(f'<span class="src-chip">{src}</span> similarity: {r["similarity"]:.2f}', unsafe_allow_html=True)
                st.caption(r["text"][:300] + ("..." if len(r["text"]) > 300 else ""))

        # Export buttons
        ec1, ec2 = st.columns(2)
        pdf_bytes = generate_pdf(report_payload)
        docx_bytes = generate_docx(report_payload)
        stamp = int(time.time())
        pdf_name = f"sentry_investigation_{stamp}.pdf"
        docx_name = f"sentry_investigation_{stamp}.docx"
        with ec1:
            if st.download_button("⬇️ Download Investigation Report (PDF)", data=pdf_bytes, file_name=pdf_name, mime="application/pdf"):
                save_report_file(pdf_name, pdf_bytes)
        with ec2:
            if st.download_button("⬇️ Download Investigation Report (DOCX)", data=docx_bytes, file_name=docx_name,
                                   mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"):
                save_report_file(docx_name, docx_bytes)
        # Always persist a copy so it shows up under Reports / History regardless of click
        save_report_file(pdf_name, pdf_bytes)
        save_report_file(docx_name, docx_bytes)

        log_investigation(
            {
                "question": query,
                "answer": answer,
                "mode": st.session_state.explanation_mode,
                "cve_ids": list(risk_data.keys()) if risk_data else [],
                "confidence": confidence.score,
                "confidence_label": confidence.label,
                "risk": {k: v.score for k, v in risk_data.items()},
                "priority": {k: v.priority for k, v in risk_data.items()},
                "pdf_report": pdf_name,
                "docx_report": docx_name,
            }
        )

    st.session_state.messages.append({"role": "assistant", "content": answer})


# =============================================================================
# PAGE: Upload Report
# =============================================================================
def page_upload():
    st.markdown('<p class="console-label">📄 Upload Vulnerability Report</p>', unsafe_allow_html=True)
    st.caption("Supported formats: PDF, DOCX, TXT, CSV, XLSX, JSON.")

    files = st.file_uploader(
        "Drop one or more report files", type=["pdf", "docx", "txt", "csv", "xlsx", "json"], accept_multiple_files=True
    )
    if files and st.button("Ingest into knowledge base"):
        for f in files:
            with st.spinner(f"Processing {f.name}..."):
                try:
                    text = extract_text(f.read(), f.name)
                    chunks = chunk_text(text)
                    n = store.add_report_chunks(chunks, f.name)
                    st.success(f"✅ Added **{n}** chunks from `{f.name}`")
                except Exception as exc:
                    st.error(f"Failed to process `{f.name}`: {exc}")

    st.markdown('<p class="console-label">📚 Indexed Reports</p>', unsafe_allow_html=True)
    sources = store.report_sources()
    if not sources:
        st.info("No reports uploaded yet.")
    else:
        for s in sources:
            st.markdown(f'<span class="src-chip">{s}</span>', unsafe_allow_html=True)

    st.markdown('<p class="console-label">⚡ Quick Actions</p>', unsafe_allow_html=True)
    if sources:
        qa1, qa2, qa3 = st.columns(3)
        if qa1.button("Summarize latest report"):
            st.session_state.messages.append({"role": "user", "content": "Summarize this report."})
            st.session_state.nav = "Investigate (Chat)"
            st.rerun()
        if qa2.button("Extract all CVEs"):
            st.session_state.messages.append({"role": "user", "content": "Extract all CVEs mentioned in this report."})
            st.session_state.nav = "Investigate (Chat)"
            st.rerun()
        if qa3.button("Generate executive summary"):
            st.session_state.messages.append({"role": "user", "content": "Generate an executive summary of this report."})
            st.session_state.nav = "Investigate (Chat)"
            st.rerun()


# =============================================================================
# PAGE: Investigation History
# =============================================================================
def page_history():
    st.markdown('<p class="console-label">🕵 Investigation History</p>', unsafe_allow_html=True)
    history = load_history()
    if not history:
        st.info("No investigations logged yet.")
        return
    for h in history:
        ts = pd.to_datetime(h["timestamp"], unit="s").strftime("%Y-%m-%d %H:%M")
        with st.expander(f"[{ts}] {h['question']}"):
            st.markdown(h["answer"])
            if h.get("cve_ids"):
                for cid in h["cve_ids"]:
                    priority = h.get("priority", {}).get(cid, "N/A")
                    score = h.get("risk", {}).get(cid, "N/A")
                    st.markdown(sev_chip(priority, f"{score}/100" if score != "N/A" else ""), unsafe_allow_html=True)
            st.caption(f"Confidence: {h.get('confidence', 'N/A')}% ({h.get('confidence_label', 'N/A')}) · Mode: {h.get('mode', 'N/A')}")


# =============================================================================
# PAGE: Reports
# =============================================================================
def page_reports():
    st.markdown('<p class="console-label">📑 Generated Reports</p>', unsafe_allow_html=True)
    files = list_report_files()
    if not files:
        st.info("No reports generated yet. Run an investigation and export a report to see it here.")
        return
    for path in files:
        col1, col2 = st.columns([4, 1])
        col1.markdown(f"**{path.name}**  \n<span style='color:var(--muted);font-size:0.8rem;'>{time.ctime(path.stat().st_mtime)}</span>", unsafe_allow_html=True)
        mime = "application/pdf" if path.suffix == ".pdf" else "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        col2.download_button("⬇️ Download", data=path.read_bytes(), file_name=path.name, mime=mime, key=f"dl-{path.name}")


# =============================================================================
# PAGE: Settings
# =============================================================================
def page_settings():
    st.markdown('<p class="console-label">⚙️ Settings</p>', unsafe_allow_html=True)

    st.markdown("**LLM Provider**")
    provider = st.selectbox("Preferred provider", ["auto", "gemini", "groq"], index=["auto", "gemini", "groq"].index(settings.llm_provider))
    gemini_key = st.text_input("Gemini API Key", type="password", value=settings.gemini_api_key or "")
    groq_key = st.text_input("Groq API Key", type="password", value=settings.groq_api_key or "")
    nvd_key = st.text_input("NVD API Key (optional — raises rate limits)", type="password", value=settings.nvd_api_key or "")

    if st.button("Save settings"):
        settings.llm_provider = provider
        settings.gemini_api_key = gemini_key or None
        settings.groq_api_key = groq_key or None
        settings.nvd_api_key = nvd_key or None
        reset_llm_client()
        st.success("Settings updated for this session. For persistence across restarts, set these as environment "
                    "variables or in `.streamlit/secrets.toml`.")

    st.divider()
    st.markdown("**Knowledge Base**")
    if st.button("🗑️ Clear uploaded reports"):
        get_store().report_col.delete(where={"source": {"$ne": "__never__"}}) if store.report_col.count() else None
        st.success("Cleared uploaded report chunks (CVE cache preserved).")


# ---------------------------------------------------------------------------
# Router
# ---------------------------------------------------------------------------
PAGES = {
    "Dashboard": page_dashboard,
    "Investigate (Chat)": page_chat,
    "Upload Report": page_upload,
    "Investigation History": page_history,
    "Reports": page_reports,
    "Settings": page_settings,
}
PAGES[st.session_state.nav]()
