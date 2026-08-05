"""
core/ingestion.py — Multi-format report ingestion.

Supports: PDF, DOCX, TXT, CSV, XLSX, JSON.
Pipeline: extract_text() -> chunk_text() -> (caller embeds + stores in Chroma)
"""
from __future__ import annotations

import io
import json
import logging
import re

import pandas as pd

from config import settings

logger = logging.getLogger("sentry.ingestion")

SUPPORTED_EXTENSIONS = {"pdf", "docx", "txt", "csv", "xlsx", "json"}


def extract_text(file_bytes: bytes, filename: str) -> str:
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    if ext not in SUPPORTED_EXTENSIONS:
        raise ValueError(f"Unsupported file type: .{ext}. Supported: {sorted(SUPPORTED_EXTENSIONS)}")

    if ext == "pdf":
        return _extract_pdf(file_bytes)
    if ext == "docx":
        return _extract_docx(file_bytes)
    if ext == "txt":
        return file_bytes.decode("utf-8", errors="ignore")
    if ext == "csv":
        return _extract_tabular(file_bytes, fmt="csv")
    if ext == "xlsx":
        return _extract_tabular(file_bytes, fmt="xlsx")
    if ext == "json":
        return _extract_json(file_bytes)
    raise ValueError(f"Unhandled extension: {ext}")  # pragma: no cover


def _extract_pdf(file_bytes: bytes) -> str:
    import fitz  # PyMuPDF

    doc = fitz.open(stream=file_bytes, filetype="pdf")
    return "\n".join(page.get_text() for page in doc)


def _extract_docx(file_bytes: bytes) -> str:
    import docx  # python-docx

    document = docx.Document(io.BytesIO(file_bytes))
    parts = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts)


def _extract_tabular(file_bytes: bytes, fmt: str) -> str:
    if fmt == "csv":
        df = pd.read_csv(io.BytesIO(file_bytes))
    else:
        df = pd.read_excel(io.BytesIO(file_bytes))
    lines = [" | ".join(str(c) for c in df.columns)]
    for _, row in df.iterrows():
        lines.append(" | ".join(str(v) for v in row.values))
    return "\n".join(lines)


def _extract_json(file_bytes: bytes) -> str:
    data = json.loads(file_bytes.decode("utf-8", errors="ignore"))
    return json.dumps(data, indent=2)


def chunk_text(text: str, size: int | None = None, overlap: int | None = None) -> list[str]:
    size = size or settings.chunk_size
    overlap = overlap or settings.chunk_overlap
    text = re.sub(r"\s+", " ", text).strip()
    chunks = []
    start = 0
    while start < len(text):
        end = start + size
        chunks.append(text[start:end])
        start += size - overlap
    return [c for c in chunks if len(c.strip()) > 30]
